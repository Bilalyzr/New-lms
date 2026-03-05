"""
Hexoria Academy - Word Document Generator
Generates a professional .docx documentation file
Run: pip install python-docx && python generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Brand Colors ────────────────────────────────────────────────────────────
PURPLE      = RGBColor(0x6C, 0x4C, 0xF1)   # #6C4CF1
DARK_PURPLE = RGBColor(0x4B, 0x3A, 0xC7)   # #4B3AC7
TEAL        = RGBColor(0x00, 0xC2, 0xA8)   # #00C2A8
DARK        = RGBColor(0x1E, 0x29, 0x3B)   # #1E293B
MUTED       = RGBColor(0x64, 0x74, 0x8B)   # #64748B
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG    = RGBColor(0xED, 0xE9, 0xFE)   # #EDE9FE (light purple)
TABLE_HDR   = RGBColor(0x6C, 0x4C, 0xF1)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, border_color="6C4CF1"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, header_bg="6C4CF1"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "F8F7FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = val
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK

    return table

def add_cover_page(doc):
    doc.add_picture  # placeholder; we'll use shapes
    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("HEXORIA ACADEMY")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = PURPLE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Technical Documentation")
    run2.font.size = Pt(20)
    run2.font.color.rgb = DARK_PURPLE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Learning Management System — v1.0")
    run3.font.size = Pt(12)
    run3.font.color.rgb = MUTED

    doc.add_paragraph()
    doc.add_paragraph()

    # Info table
    info = [
        ("Project Name", "Hexoria Academy LMS"),
        ("Type", "Cloud-Native SaaS LMS"),
        ("Frontend", "React 18 + Vite"),
        ("Backend", "Node.js + Express"),
        ("Database", "PostgreSQL (AWS RDS)"),
        ("Cloud", "AWS (EC2, S3, CloudFront, RDS)"),
        ("Version", "1.0 — Phase 8 Complete"),
        ("Date", datetime.date.today().strftime("%B %d, %Y")),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_bg(row.cells[0], "EDE9FE")
        r0 = row.cells[0].paragraphs[0].runs[0]
        r0.font.bold = True
        r0.font.color.rgb = PURPLE
        r0.font.size = Pt(10)
        r1 = row.cells[1].paragraphs[0].runs[0]
        r1.font.size = Pt(10)
        r1.font.color.rgb = DARK
    doc.add_page_break()

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = PURPLE
        p.runs[0].font.size = Pt(18)
    elif level == 2:
        p.runs[0].font.color.rgb = DARK_PURPLE
        p.runs[0].font.size = Pt(14)
    elif level == 3:
        p.runs[0].font.color.rgb = DARK
        p.runs[0].font.size = Pt(12)
    return p

def add_para(doc, text, color=None, size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color if color else DARK
    return p

def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color if color else DARK

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x31, 0x48)

def add_section_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '6C4CF1')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─── Document Builder ─────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default paragraph font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── COVER PAGE ──────────────────────────────────────────────────────────
    add_cover_page(doc)

    # ── TABLE OF CONTENTS ────────────────────────────────────────────────────
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        ("1.", "System Architecture"),
        ("2.", "System Feasibility Analysis"),
        ("3.", "Use Case Diagrams"),
        ("4.", "Sequence Diagrams"),
        ("5.", "Entity Relationship Diagram"),
        ("6.", "Module Descriptions"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{num}  ")
        r1.font.bold = True
        r1.font.color.rgb = PURPLE
        r1.font.size = Pt(11)
        r2 = p.add_run(title)
        r2.font.color.rgb = DARK
        r2.font.size = Pt(11)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1: SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. System Architecture", 1)
    add_section_divider(doc)

    add_heading(doc, "1.1 Overview", 2)
    add_para(doc,
        "Hexoria Academy is a cloud-native, multi-tier SaaS LMS built using a decoupled "
        "architecture where the frontend, backend, and storage layers operate independently "
        "and communicate through well-defined REST APIs.")

    add_heading(doc, "1.2 Architecture Layers", 2)

    add_heading(doc, "Layer 1 — Presentation Layer (Frontend)", 3)
    add_styled_table(doc,
        ["Component", "Technology", "Purpose"],
        [
            ("UI Framework",       "React 18 + Vite",       "SPA rendering"),
            ("Routing",            "React Router v6",       "Client-side navigation"),
            ("State Management",   "React Context API",     "Auth state, global state"),
            ("HTTP Client",        "Axios",                 "API communication"),
            ("Styling",            "Vanilla CSS + Variables","Brand theming"),
            ("Icons",              "Lucide React",          "UI icon set"),
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "Frontend Route Map", 3)
    add_styled_table(doc,
        ["Route", "Page", "Access Level"],
        [
            ("/",                         "Home (Landing)",       "Public"),
            ("/courses",                  "Course Listing",       "Public"),
            ("/course/:id",               "Course Details",       "Public"),
            ("/dashboard",                "Student Dashboard",    "Student"),
            ("/learn/:courseId/:lessonId","Course Player",        "Student (Enrolled)"),
            ("/cart",                     "Shopping Cart",        "Student"),
            ("/wishlist",                 "Wishlist",             "Student"),
            ("/instructor",               "Instructor Dashboard", "Instructor"),
            ("/build",                    "Course Builder",       "Instructor"),
            ("/admin",                    "Admin Dashboard",      "Admin"),
            ("/settings",                 "Settings",             "All Roles"),
            ("/notifications",            "Notifications",        "All Roles"),
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "Layer 2 — Application Layer (Backend)", 3)
    add_styled_table(doc,
        ["Component", "Technology", "Purpose"],
        [
            ("Runtime",         "Node.js",          "Server-side JavaScript"),
            ("Framework",       "Express.js",       "REST API routing"),
            ("Authentication",  "JWT",              "Stateless auth tokens"),
            ("Password Security","bcrypt",          "Password hashing"),
            ("DB Driver",       "node-postgres(pg)","PostgreSQL connection pool"),
            ("Config",          "dotenv",           "Environment variables"),
            ("CORS",            "cors middleware",  "Cross-origin request handling"),
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "API Route Structure", 3)
    add_code_block(doc,
        "/api/auth\n"
        "  ├── POST /register   → Create new user account\n"
        "  └── POST /login      → Authenticate & return JWT\n\n"
        "/api/courses\n"
        "  ├── GET  /           → Fetch all published courses\n"
        "  ├── GET  /:id        → Fetch single course details\n"
        "  ├── POST /           → Create new course (Instructor)\n"
        "  └── PUT  /:id        → Update course (Instructor)"
    )
    doc.add_paragraph()

    add_heading(doc, "Layer 3 — Data Layer (PostgreSQL)", 3)
    add_styled_table(doc,
        ["Table", "Primary Fields", "Purpose"],
        [
            ("users",       "id, full_name, email, password_hash, role",        "All platform users"),
            ("courses",     "id, title, category, price, instructor_id, status","Course catalog"),
            ("sections",    "id, course_id, title, order_index",                "Curriculum structure"),
            ("lessons",     "id, section_id, title, video_url, duration",       "Individual lessons"),
            ("enrollments", "id, student_id, course_id, progress_percentage",   "Course access"),
            ("cart_items",  "id, user_id, course_id",                           "Shopping cart"),
            ("wishlist",    "id, user_id, course_id",                           "Saved courses"),
            ("payments",    "id, user_id, course_id, amount, status",           "Payment records"),
            ("notifications","id, user_id, title, message, is_read",            "In-app alerts"),
            ("certificates","id, student_id, course_id, issued_at",             "Completion certs"),
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "Layer 4 — Cloud Infrastructure (AWS)", 3)
    aws_items = [
        "CloudFront CDN — Global content delivery + static asset caching",
        "S3 (Frontend) — Hosts React production build files",
        "EC2 (t3.small) — Runs Node.js Express API via PM2 + Nginx reverse proxy",
        "RDS PostgreSQL — Managed relational database with automated backups",
        "S3 (Media) — Private bucket for video storage with pre-signed URLs",
        "ACM — SSL/TLS certificate management",
        "Route 53 — DNS routing and domain management",
    ]
    for item in aws_items:
        add_bullet(doc, item)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2: SYSTEM FEASIBILITY
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. System Feasibility Analysis", 1)
    add_section_divider(doc)

    add_heading(doc, "2.1 Technical Feasibility", 2)
    add_styled_table(doc,
        ["Factor", "Assessment"],
        [
            ("Tech Stack Maturity",  "React, Node.js, PostgreSQL are industry-proven technologies"),
            ("Team Skill Match",     "JavaScript end-to-end — minimal context switching required"),
            ("Cloud Readiness",      "AWS services are production-grade and widely adopted"),
            ("Scalability",          "Stateless API + managed RDS + CDN enable horizontal scaling"),
            ("Security",             "JWT, bcrypt, HTTPS, private S3 buckets, role-based middleware"),
            ("Development Tools",    "Vite (fast builds), dotenv (config), pg Pool (DB connections)"),
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "Technical Risks & Mitigations", 3)
    add_styled_table(doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ("Video overload on EC2",  "High cost",     "Stream via S3 + CloudFront signed URLs"),
            ("JWT token theft",        "Security breach","Short expiry + HTTPS-only communication"),
            ("DB connection overflow", "API crashes",   "pg Pool + RDS max connection limits"),
            ("Stripe webhook failure", "Missed payments","Webhook signature + idempotency keys"),
            ("Cold start latency",     "Poor first UX", "CloudFront caching + React code splitting"),
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "2.2 Operational Feasibility", 2)
    add_styled_table(doc,
        ["Role", "Key Operations", "Status"],
        [
            ("Student",    "Browse, Purchase, Learn, Track progress",          "✅ Fully Implemented"),
            ("Instructor", "Create courses, Upload content, Track revenue",    "✅ Course Builder Done"),
            ("Admin",      "Manage users, Approve courses, Monitor payments",  "✅ Dashboard Done"),
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "2.3 Economic Feasibility", 2)
    add_styled_table(doc,
        ["AWS Service", "Tier", "Est. Cost/Month"],
        [
            ("EC2 (t3.small)",          "Backend API server",   "~$15–20"),
            ("RDS PostgreSQL",           "db.t3.micro",          "~$15–25"),
            ("S3 Storage (50GB videos)", "Media storage",        "~$1–5"),
            ("CloudFront CDN",           "Global delivery",      "~$5–15"),
            ("Route 53 + ACM",           "DNS + SSL",            "~$1–2"),
            ("Total Estimate",           "—",                    "~$37–67/month"),
        ]
    )
    doc.add_paragraph()
    add_para(doc,
        "Revenue Model: At 10% commission on $29/course, only ~150 course sales/month "
        "covers full infrastructure. Revenue scales exponentially with student growth.",
        color=MUTED, italic=True)

    add_heading(doc, "2.4 Schedule Feasibility", 2)
    add_styled_table(doc,
        ["Phase", "Description", "Status"],
        [
            ("Phase 1", "Core Foundation — Setup, Home, Navbar",                  "✅ Complete"),
            ("Phase 2", "Role Dashboards — Student, Instructor, Admin",            "✅ Complete"),
            ("Phase 3", "Authentication — Login, Register, JWT, AuthContext",      "✅ Complete"),
            ("Phase 4", "E-Commerce — Cart, Wishlist, Checkout, CourseDetails",    "✅ Complete"),
            ("Phase 5", "Learning System — Course Player, Progress Tracking",      "✅ Complete"),
            ("Phase 6", "Settings, Notifications, Search Integration",             "✅ Complete"),
            ("Phase 7", "Live Backend — PostgreSQL Schema, Auth API, Courses API", "✅ Complete"),
            ("Phase 8", "Production Build — Full-stack local test + React build",  "✅ Complete"),
            ("Phase 9", "AWS Deployment — RDS, EC2, S3, CloudFront, Domain",       "🚀 In Progress"),
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3: USE CASE DIAGRAMS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Use Case Diagrams", 1)
    add_section_divider(doc)
    add_para(doc,
        "Note: The following use cases are described textually. For visual UML diagrams, "
        "refer to the Mermaid-rendered documentation or submit these descriptions to a UML tool.",
        color=MUTED, italic=True, size=10)
    doc.add_paragraph()

    add_heading(doc, "3.1 Student Use Cases", 2)
    student_ucs = [
        "UC-S1:  Browse the public course catalog",
        "UC-S2:  Search and filter courses by keyword, category, level, price",
        "UC-S3:  View full course details page",
        "UC-S4:  Register a new student account",
        "UC-S5:  Login with email and password",
        "UC-S6:  Add a course to Wishlist",
        "UC-S7:  Add a course to Shopping Cart",
        "UC-S8:  Checkout and pay via Stripe",
        "UC-S9:  Access enrolled courses after successful payment",
        "UC-S10: Watch video lessons in the Course Player",
        "UC-S11: Mark a lesson as complete",
        "UC-S12: Track overall course progress percentage",
        "UC-S13: Download completion certificate",
        "UC-S14: Edit profile and change password in Settings",
        "UC-S15: View in-app notifications",
    ]
    for uc in student_ucs:
        add_bullet(doc, uc)
    doc.add_paragraph()

    add_heading(doc, "3.2 Instructor Use Cases", 2)
    inst_ucs = [
        "UC-I1:  Register as an instructor",
        "UC-I2:  Login to instructor dashboard",
        "UC-I3:  Create a new course via multi-step Course Builder",
        "UC-I4:  Fill in course basics (title, description, price, difficulty)",
        "UC-I5:  Build curriculum with sections and lessons",
        "UC-I6:  Upload video lessons to AWS S3",
        "UC-I7:  Set additional info (requirements, audience, FAQs, SEO)",
        "UC-I8:  Publish course or save as draft",
        "UC-I9:  Edit existing published courses",
        "UC-I10: View list of enrolled students",
        "UC-I11: Track total revenue and earnings",
        "UC-I12: Manage instructor bio and payout details",
        "UC-I13: Receive notifications for new enrollments",
    ]
    for uc in inst_ucs:
        add_bullet(doc, uc)
    doc.add_paragraph()

    add_heading(doc, "3.3 Admin Use Cases", 2)
    admin_ucs = [
        "UC-A1:  Login to the admin dashboard",
        "UC-A2:  View platform-wide analytics (users, courses, revenue)",
        "UC-A3:  Manage all user accounts (list, search, filter)",
        "UC-A4:  Approve or remove courses submitted by instructors",
        "UC-A5:  Monitor all payment transactions",
        "UC-A6:  Set the platform commission rate",
        "UC-A7:  Configure global platform settings",
        "UC-A8:  Generate and view system reports",
        "UC-A9:  Send platform-wide notifications",
        "UC-A10: Suspend or activate any user account",
    ]
    for uc in admin_ucs:
        add_bullet(doc, uc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4: SEQUENCE DIAGRAMS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Sequence Diagrams", 1)
    add_section_divider(doc)
    add_para(doc,
        "The following sequence diagrams describe the step-by-step message flows "
        "between actors and system components for the key platform operations.",
        color=MUTED, size=10, italic=True)
    doc.add_paragraph()

    sequences = [
        {
            "title": "4.1 User Login Flow",
            "steps": [
                ("User",    "React",   "Enter email & password → Submit Login form"),
                ("React",   "API",     "POST /api/auth/login { email, password }"),
                ("API",     "DB",      "SELECT * FROM users WHERE email = ?"),
                ("DB",      "API",     "Return user row (with password_hash)"),
                ("API",     "API",     "bcrypt.compare(password, hash)"),
                ("API",     "React",   "[IF VALID] 200 OK + JWT token + user object"),
                ("API",     "React",   "[IF INVALID] 401 Unauthorized"),
                ("React",   "User",    "Redirect to role dashboard OR show error"),
            ]
        },
        {
            "title": "4.2 Student Course Enrollment Flow",
            "steps": [
                ("Student", "React",  "Click 'Enroll Now' on Course Details page"),
                ("React",   "API",    "POST /api/cart/:courseId (JWT Bearer token)"),
                ("API",     "API",    "Verify JWT token + extract user ID"),
                ("API",     "DB",     "INSERT INTO cart_items (user_id, course_id)"),
                ("DB",      "API",    "201 Created"),
                ("Student", "React",  "Click 'Checkout'"),
                ("React",   "Stripe", "Create Payment Intent → get client_secret"),
                ("Student", "Stripe", "Enter card details & confirm payment"),
                ("Stripe",  "React",  "Payment successful event"),
                ("React",   "API",    "POST /api/enrollments/:courseId (JWT)"),
                ("API",     "DB",     "INSERT INTO enrollments (student_id, course_id)"),
                ("API",     "DB",     "DELETE FROM cart_items WHERE course_id = ?"),
                ("API",     "React",  "201 Enrolled successfully"),
                ("React",   "Student","Redirect to Course Player ✅"),
            ]
        },
        {
            "title": "4.3 Instructor Course Creation Flow",
            "steps": [
                ("Instructor", "React", "Open Course Builder wizard"),
                ("Instructor", "React", "Step 1: Fill title, description, price, difficulty"),
                ("Instructor", "React", "Step 2: Add sections, add lessons, upload video"),
                ("React",      "S3",    "PUT video file to private S3 bucket"),
                ("S3",         "React", "Return public video_url"),
                ("Instructor", "React", "Step 3: Fill requirements, FAQs, SEO metadata"),
                ("Instructor", "React", "Step 4: Click 'Publish'"),
                ("React",      "API",   "POST /api/courses { title, sections, lessons, video_url } (JWT)"),
                ("API",        "API",   "Verify JWT + confirm role === 'instructor'"),
                ("API",        "DB",    "INSERT INTO courses (status='Published')"),
                ("API",        "DB",    "INSERT INTO sections and lessons"),
                ("DB",         "API",   "Success"),
                ("API",        "React", "201 Course Created"),
                ("React",      "Instructor", "Show success toast notification 🎉"),
            ]
        },
        {
            "title": "4.4 Lesson Completion Flow",
            "steps": [
                ("Student", "Player", "Open enrolled course in Course Player"),
                ("Player",  "API",   "GET /api/courses/:id/lessons (JWT)"),
                ("API",     "DB",    "SELECT lessons WHERE course_id = ?"),
                ("DB",      "API",   "Return ordered lesson list"),
                ("API",     "Player","Lesson data with video URLs"),
                ("Player",  "Student","Display video player + lesson sidebar"),
                ("Student", "Player", "Watch video → click 'Mark as Complete'"),
                ("Player",  "API",   "POST /api/progress/:lessonId (JWT)"),
                ("API",     "DB",    "INSERT INTO progress (student_id, lesson_id, completed=true)"),
                ("API",     "DB",    "UPDATE enrollments SET progress_percentage = calculated"),
                ("DB",      "API",   "Updated"),
                ("API",     "Player","{ progress: 75 }"),
                ("Player",  "Student","Progress bar updates ✅"),
                ("API",     "DB",    "[IF 100%] INSERT INTO certificates (student_id, course_id)"),
                ("Player",  "Student","[IF 100%] Show 'Download Certificate' button 🎓"),
            ]
        },
    ]

    for seq in sequences:
        add_heading(doc, seq["title"], 2)
        table = doc.add_table(rows=1 + len(seq["steps"]), cols=4)
        table.style = 'Table Grid'

        # Header
        heads = ["#", "From", "To", "Message / Action"]
        hrow = table.rows[0]
        for i, h in enumerate(heads):
            hrow.cells[i].text = h
            set_cell_bg(hrow.cells[i], "6C4CF1")
            run = hrow.cells[i].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)

        for si, (frm, to, msg) in enumerate(seq["steps"]):
            row = table.rows[si + 1]
            bg = "F8F7FF" if si % 2 == 0 else "FFFFFF"
            vals = [str(si + 1), frm, to, msg]
            for ci, val in enumerate(vals):
                row.cells[ci].text = val
                set_cell_bg(row.cells[ci], bg)
                run = row.cells[ci].paragraphs[0].runs[0]
                run.font.size = Pt(9)
                run.font.color.rgb = PURPLE if ci in [1, 2] else DARK
                if ci in [1, 2]:
                    run.font.bold = True

        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5: ER DIAGRAM  (textual — Mermaid in markdown version)
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Entity Relationship Diagram", 1)
    add_section_divider(doc)
    add_para(doc,
        "The following describes the complete database schema for Hexoria Academy. "
        "All entities, their attributes, and relationships are defined below.",
        color=MUTED, size=10, italic=True)
    doc.add_paragraph()

    entities = [
        {
            "name": "USERS",
            "fields": [
                ("id",            "SERIAL",       "PK", "Primary key"),
                ("full_name",     "VARCHAR(100)", "",   "User's full name"),
                ("email",         "VARCHAR(100)", "UK", "Unique email address"),
                ("password_hash", "VARCHAR(255)", "",   "bcrypt hashed password"),
                ("role",          "VARCHAR(20)",  "",   "student | instructor | admin"),
                ("created_at",    "TIMESTAMP",    "",   "Registration timestamp"),
            ]
        },
        {
            "name": "COURSES",
            "fields": [
                ("id",               "SERIAL",        "PK", "Primary key"),
                ("title",            "VARCHAR(255)",   "",   "Course title"),
                ("category",         "VARCHAR(50)",    "",   "Subject category"),
                ("price",            "DECIMAL(10,2)",  "",   "Course price in USD"),
                ("description",      "TEXT",           "",   "Full course description"),
                ("instructor_id",    "INT",            "FK", "References users(id)"),
                ("image_url",        "TEXT",           "",   "Thumbnail image URL"),
                ("hours",            "INT",            "",   "Total course duration"),
                ("rating",           "DECIMAL(3,2)",   "",   "Average student rating"),
                ("students_enrolled","INT",            "",   "Total enrolled count"),
                ("status",           "VARCHAR(20)",    "",   "Draft | Published"),
                ("created_at",       "TIMESTAMP",      "",   "Creation timestamp"),
            ]
        },
        {
            "name": "SECTIONS",
            "fields": [
                ("id",          "SERIAL",       "PK", "Primary key"),
                ("course_id",   "INT",          "FK", "References courses(id)"),
                ("title",       "VARCHAR(255)", "",   "Section heading"),
                ("order_index", "INT",          "",   "Display order"),
            ]
        },
        {
            "name": "LESSONS",
            "fields": [
                ("id",          "SERIAL",       "PK", "Primary key"),
                ("section_id",  "INT",          "FK", "References sections(id)"),
                ("title",       "VARCHAR(255)", "",   "Lesson title"),
                ("video_url",   "TEXT",         "",   "S3 video URL"),
                ("duration",    "VARCHAR(20)",  "",   "Lesson duration"),
                ("order_index", "INT",          "",   "Display order"),
            ]
        },
        {
            "name": "ENROLLMENTS",
            "fields": [
                ("id",                  "SERIAL",    "PK",  "Primary key"),
                ("student_id",          "INT",       "FK",  "References users(id)"),
                ("course_id",           "INT",       "FK",  "References courses(id)"),
                ("progress_percentage", "INT",       "",    "0–100% completion"),
                ("enrolled_at",         "TIMESTAMP", "",    "Enrollment timestamp"),
            ]
        },
        {
            "name": "CART_ITEMS",
            "fields": [
                ("id",        "SERIAL",    "PK", "Primary key"),
                ("user_id",   "INT",       "FK", "References users(id)"),
                ("course_id", "INT",       "FK", "References courses(id)"),
                ("added_at",  "TIMESTAMP", "",   "When added to cart"),
            ]
        },
        {
            "name": "PAYMENTS",
            "fields": [
                ("id",        "SERIAL",       "PK", "Primary key"),
                ("user_id",   "INT",          "FK", "References users(id)"),
                ("course_id", "INT",          "FK", "References courses(id)"),
                ("amount",    "DECIMAL(10,2)","",   "Payment amount in USD"),
                ("status",    "VARCHAR(20)",  "",   "pending | completed | failed"),
                ("stripe_id", "VARCHAR(255)", "",   "Stripe payment intent ID"),
                ("paid_at",   "TIMESTAMP",   "",   "Payment timestamp"),
            ]
        },
        {
            "name": "NOTIFICATIONS",
            "fields": [
                ("id",         "SERIAL",    "PK", "Primary key"),
                ("user_id",    "INT",       "FK", "References users(id)"),
                ("title",      "VARCHAR",   "",   "Notification title"),
                ("message",    "TEXT",      "",   "Notification body"),
                ("is_read",    "BOOLEAN",   "",   "Read/unread flag"),
                ("created_at", "TIMESTAMP", "",   "Alert timestamp"),
            ]
        },
        {
            "name": "CERTIFICATES",
            "fields": [
                ("id",         "SERIAL",    "PK", "Primary key"),
                ("student_id", "INT",       "FK", "References users(id)"),
                ("course_id",  "INT",       "FK", "References courses(id)"),
                ("issued_at",  "TIMESTAMP", "",   "Certificate issue date"),
            ]
        },
    ]

    for entity in entities:
        add_heading(doc, f"Table: {entity['name']}", 3)
        add_styled_table(doc,
            ["Column", "Type", "Key", "Description"],
            entity["fields"]
        )
        doc.add_paragraph()

    add_heading(doc, "Relationships Summary", 2)
    relationships = [
        ("USERS",       "COURSES",       "One-to-Many",  "One instructor creates many courses"),
        ("USERS",       "ENROLLMENTS",   "One-to-Many",  "One student has many enrollments"),
        ("COURSES",     "ENROLLMENTS",   "One-to-Many",  "One course has many enrolled students"),
        ("COURSES",     "SECTIONS",      "One-to-Many",  "One course contains many sections"),
        ("SECTIONS",    "LESSONS",       "One-to-Many",  "One section contains many lessons"),
        ("USERS",       "CART_ITEMS",    "One-to-Many",  "One user has many cart items"),
        ("USERS",       "PAYMENTS",      "One-to-Many",  "One user makes many payments"),
        ("USERS",       "NOTIFICATIONS", "One-to-Many",  "One user receives many notifications"),
        ("USERS",       "CERTIFICATES",  "One-to-Many",  "One student earns many certificates"),
        ("COURSES",     "CERTIFICATES",  "One-to-Many",  "One course awards many certificates"),
    ]
    add_styled_table(doc,
        ["From Table", "To Table", "Relationship", "Description"],
        relationships
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6: MODULE DESCRIPTIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Module Descriptions", 1)
    add_section_divider(doc)

    modules = [
        {
            "num": "Module 1",
            "name": "Authentication & Authorization",
            "path": "backend/src/routes/auth.js · frontend/src/context/AuthContext.jsx",
            "purpose": "Manages all user identity operations — registration, login, session persistence, and route protection across all three roles.",
            "features": [
                "POST /api/auth/register — Hashes password with bcrypt, inserts user, returns JWT",
                "POST /api/auth/login — Validates credentials, issues signed JWT with role payload",
                "AuthContext — Global React context exposing user, login(), logout() to all components",
                "authMiddleware.js — Verifies Bearer JWT on all protected routes",
                "Role-based access: student | instructor | admin enforced at API layer",
            ]
        },
        {
            "num": "Module 2",
            "name": "Course Catalog & Listing",
            "path": "frontend/src/pages/CourseListing.jsx · backend/src/routes/courses.js",
            "purpose": "Allows all users to browse and filter the full course catalog, with real-time search.",
            "features": [
                "Keyword search via query params (GET /api/courses?q=react)",
                "Filter by category, difficulty, and price range",
                "Course cards with title, rating, instructor, price, student count",
                "Animated card hover effects (lift + shadow increase)",
                "Responsive CSS grid layout for all screen sizes",
            ]
        },
        {
            "num": "Module 3",
            "name": "Course Details Page",
            "path": "frontend/src/pages/CourseDetails.jsx",
            "purpose": "Full product page for a course with all information needed to support purchase decisions.",
            "features": [
                "Course title, description, instructor profile",
                "Expandable curriculum accordion (sections → lessons)",
                "Requirements, Target Audience, Learning Outcomes sections",
                "Enroll, Add to Cart, Add to Wishlist CTAs",
                "Price color: Free = teal (#00C2A8), Paid = purple (#6C4CF1)",
            ]
        },
        {
            "num": "Module 4",
            "name": "Course Builder (Multi-Step Wizard)",
            "path": "frontend/src/pages/CourseBuilder.jsx",
            "purpose": "Enables instructors to create complete, structured courses through a guided 4-step wizard.",
            "features": [
                "Step 1 — Basics: Title, Description, Thumbnail, Intro Video, Price, Difficulty",
                "Step 2 — Curriculum: Add Sections → Add Lessons → Upload Videos to S3",
                "Step 3 — Additional: Requirements, FAQs, Tags, SEO Metadata, Duration",
                "Step 4 — Review & Publish: Validate fields → Save as Draft or Publish",
                "Progress indicator showing current wizard step",
            ]
        },
        {
            "num": "Module 5",
            "name": "Course Player (Learning System)",
            "path": "frontend/src/pages/CoursePlayer.jsx",
            "purpose": "Immersive dark-mode learning environment for video lessons and progress tracking.",
            "features": [
                "Dark UI (#0F172A background) for distraction-free viewing",
                "Left sidebar: ordered lesson list with completion checkmarks",
                "Video player streaming from S3 signed URLs",
                "'Mark as Complete' updates progress_percentage in enrollments",
                "Animated progress bar; certificate unlocked at 100%",
            ]
        },
        {
            "num": "Module 6",
            "name": "E-Commerce System",
            "path": "Cart.jsx · Wishlist.jsx · CourseDetails.jsx (Enroll CTA)",
            "purpose": "Handles the full shopping lifecycle from interest to enrollment.",
            "features": [
                "Wishlist: Save courses with animated heart icon pop effect",
                "Cart: Add/remove courses with real-time total calculation",
                "Checkout: Stripe payment form with animated processing and success state",
                "Enrollment created automatically after successful Stripe payment",
                "Cart items cleared on successful enrollment",
            ]
        },
        {
            "num": "Module 7",
            "name": "Student Dashboard",
            "path": "frontend/src/pages/Dashboard.jsx",
            "purpose": "Personal learning hub for students to view stats and continue enrolled courses.",
            "features": [
                "Stat cards: Enrolled Courses, Completed, Hours Learned, Certificates",
                "Count-up animation on stat numbers",
                "Enrolled course list with animated purple progress bars",
                "Quick navigation links to My Courses, Certificates, Wishlist",
            ]
        },
        {
            "num": "Module 8",
            "name": "Instructor Dashboard",
            "path": "frontend/src/pages/InstructorDashboard.jsx",
            "purpose": "Business intelligence hub for instructors to manage courses and track earnings.",
            "features": [
                "Revenue counter with animated chart draw-on-load",
                "Course status: Draft (gray) → Published (purple) → Trending (teal)",
                "Enrolled student count per course",
                "Direct 'Create New Course' CTA → navigates to Course Builder",
                "Course management actions: Edit, View Students",
            ]
        },
        {
            "num": "Module 9",
            "name": "Admin Dashboard",
            "path": "frontend/src/pages/AdminDashboard.jsx",
            "purpose": "Platform-wide governance and oversight for administrator accounts.",
            "features": [
                "System-level stats: total users, total courses, total revenue",
                "User management table with activate/suspend controls",
                "Course approval and removal interface",
                "Payment and transaction monitoring",
                "Platform commission rate configuration",
            ]
        },
        {
            "num": "Module 10",
            "name": "Notifications System",
            "path": "frontend/src/pages/Notifications.jsx",
            "purpose": "In-app notification center for platform events across all user roles.",
            "features": [
                "Notification types: payment success/fail, course published, new enrollment, certificate",
                "Unread: #EDE9FE background + purple left border indicator",
                "Read: white background with fade transition on mark-as-read",
                "Notification badge with bounce animation in navbar",
                "DB: notifications(id, user_id, title, message, is_read, created_at)",
            ]
        },
        {
            "num": "Module 11",
            "name": "Settings System",
            "path": "frontend/src/pages/Settings.jsx",
            "purpose": "Role-specific user configuration panel for all authenticated users.",
            "features": [
                "Student tabs: Edit Profile, Change Password",
                "Instructor tabs: Profile, Password, Bio, Payout Details",
                "Admin tabs: Profile, Password, Platform Commission, Global Settings",
                "Sidebar tab nav with purple active highlight",
                "Save button triggers teal toast popup (slide from bottom)",
            ]
        },
        {
            "num": "Module 12",
            "name": "Database & Backend Core",
            "path": "backend/src/index.js · backend/src/config/db.js · backend/database.sql",
            "purpose": "Foundation layer powering all data persistence and API request handling.",
            "features": [
                "Express app: configures CORS, JSON parsing, mounts all route modules",
                "pg Pool: supports DATABASE_URL (production RDS) and local config via dotenv",
                "PostgreSQL schema: all tables, constraints, foreign keys, cascade deletes",
                "Auth routes: /register and /login with JWT generation",
                "Course routes: CRUD operations validated by auth middleware",
            ]
        },
    ]

    for mod in modules:
        add_heading(doc, f"{mod['num']} — {mod['name']}", 2)
        p = doc.add_paragraph()
        r = p.add_run("Path: ")
        r.font.bold = True
        r.font.color.rgb = PURPLE
        r2 = p.add_run(mod["path"])
        r2.font.name = "Courier New"
        r2.font.size = Pt(9)
        r2.font.color.rgb = MUTED

        add_para(doc, mod["purpose"], size=10.5)

        add_para(doc, "Key Features:", bold=True, color=DARK_PURPLE, size=10)
        for feat in mod["features"]:
            add_bullet(doc, feat)
        doc.add_paragraph()

    # ── FUTURE MODULES ───────────────────────────────────────────────────────
    add_heading(doc, "Future Modules (Phase 9+)", 2)
    future = [
        ("Quiz System",             "In-course assessments with auto-grading and scoring"),
        ("Certificate Generator",   "PDF certificate generation on 100% course completion"),
        ("Stripe Payouts",          "Automated instructor earnings transfers to bank accounts"),
        ("Live Classes",            "Real-time video sessions using WebRTC or third-party service"),
        ("AI Recommendations",      "ML-based course suggestions based on student learning history"),
        ("Mobile App",              "React Native cross-platform iOS and Android application"),
    ]
    add_styled_table(doc,
        ["Module", "Description"],
        future
    )

    # ── FOOTER ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hexoria Academy — Technical Documentation v1.0")
    run.font.color.rgb = PURPLE
    run.font.bold = True
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(f"Generated on {datetime.date.today().strftime('%B %d, %Y')}")
    r.font.color.rgb = MUTED
    r.font.size = Pt(10)
    r.font.italic = True

    # Save
    output_path = "Hexoria_Academy_Documentation.docx"
    doc.save(output_path)
    print(f"✅ Document saved: {output_path}")
    return output_path

if __name__ == "__main__":
    build_document()
